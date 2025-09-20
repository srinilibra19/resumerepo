"""
Document processing utilities for the Resume Chatbot application.
Handles text extraction from .txt and .docx files with validation and error handling.
"""

import io
import logging
from typing import Optional, Tuple
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Supported file formats
SUPPORTED_FORMATS = {'.txt', '.docx'}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB limit for Streamlit Cloud optimization
MEMORY_EFFICIENT_SIZE = 1 * 1024 * 1024  # 1MB threshold for memory-efficient processing

class DocumentProcessingError(Exception):
    """Custom exception for document processing errors"""
    pass

class UnsupportedFileFormatError(DocumentProcessingError):
    """Exception raised when file format is not supported"""
    pass

class FileSizeExceededError(DocumentProcessingError):
    """Exception raised when file size exceeds limit"""
    pass

class CorruptedFileError(DocumentProcessingError):
    """Exception raised when file is corrupted or unreadable"""
    pass

def validate_file_format(filename: str) -> bool:
    """
    Validate if the file format is supported.
    
    Args:
        filename (str): Name of the file to validate
        
    Returns:
        bool: True if format is supported, False otherwise
        
    Raises:
        UnsupportedFileFormatError: If file format is not supported
    """
    if not filename:
        raise UnsupportedFileFormatError("Filename cannot be empty")
    
    # Extract file extension
    file_extension = None
    if '.' in filename:
        file_extension = '.' + filename.split('.')[-1].lower()
    
    if file_extension not in SUPPORTED_FORMATS:
        supported_formats_str = ', '.join(SUPPORTED_FORMATS)
        raise UnsupportedFileFormatError(
            f"Unsupported file format '{file_extension}'. "
            f"Supported formats: {supported_formats_str}"
        )
    
    return True

def validate_file_size(file_size: int) -> bool:
    """
    Validate if the file size is within acceptable limits.
    
    Args:
        file_size (int): Size of the file in bytes
        
    Returns:
        bool: True if size is acceptable, False otherwise
        
    Raises:
        FileSizeExceededError: If file size exceeds the limit
    """
    if file_size > MAX_FILE_SIZE:
        max_size_mb = MAX_FILE_SIZE / (1024 * 1024)
        current_size_mb = file_size / (1024 * 1024)
        raise FileSizeExceededError(
            f"File size ({current_size_mb:.1f}MB) exceeds maximum allowed size ({max_size_mb}MB)"
        )
    
    return True

def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from a .txt file.
    
    Args:
        file_content (bytes): Content of the text file
        
    Returns:
        str: Extracted text content
        
    Raises:
        CorruptedFileError: If file cannot be decoded
    """
    try:
        # Try UTF-8 first, then fall back to other encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                logger.info(f"Successfully decoded text file using {encoding} encoding")
                return text.strip()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail
        raise CorruptedFileError("Unable to decode text file with any supported encoding")
        
    except Exception as e:
        logger.error(f"Error extracting text from .txt file: {str(e)}")
        raise CorruptedFileError(f"Failed to extract text from .txt file: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from a .docx file with memory optimization.
    
    Args:
        file_content (bytes): Content of the Word document
        
    Returns:
        str: Extracted text content
        
    Raises:
        CorruptedFileError: If file is corrupted or cannot be processed
    """
    try:
        # Create a file-like object from bytes
        file_stream = io.BytesIO(file_content)
        
        # Load the document
        doc = DocxDocument(file_stream)
        
        # Memory optimization: use generator for large documents
        file_size = len(file_content)
        use_memory_efficient = file_size > MEMORY_EFFICIENT_SIZE
        
        if use_memory_efficient:
            logger.info(f"Using memory-efficient processing for large document ({file_size / 1024 / 1024:.1f}MB)")
        
        # Extract text from all paragraphs
        text_content = []
        paragraph_count = 0
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_content.append(paragraph.text.strip())
                paragraph_count += 1
                
                # Memory optimization: limit processing for very large documents
                if use_memory_efficient and paragraph_count > 1000:
                    logger.warning("Large document detected, limiting paragraph processing to prevent memory issues")
                    text_content.append("\n[Document truncated for memory optimization]")
                    break
        
        # Extract text from tables if any (with memory limits)
        table_count = 0
        for table in doc.tables:
            if use_memory_efficient and table_count > 50:
                logger.warning("Large document detected, limiting table processing")
                text_content.append("\n[Additional tables truncated for memory optimization]")
                break
                
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text.strip())
            table_count += 1
        
        extracted_text = '\n'.join(text_content)
        
        if not extracted_text.strip():
            raise CorruptedFileError("Document appears to be empty or contains no readable text")
        
        logger.info(f"Successfully extracted text from .docx file ({len(extracted_text)} characters)")
        
        # Clean up memory
        del doc
        file_stream.close()
        
        return extracted_text
        
    except PackageNotFoundError:
        logger.error("Invalid .docx file format")
        raise CorruptedFileError("File is not a valid Word document or is corrupted")
    except MemoryError:
        logger.error("Memory error during .docx processing")
        raise CorruptedFileError("Document is too large to process with available memory. Please try a smaller document.")
    except Exception as e:
        logger.error(f"Error extracting text from .docx file: {str(e)}")
        raise CorruptedFileError(f"Failed to extract text from .docx file: {str(e)}")

def extract_text_from_file(file_content: bytes, filename: str) -> Tuple[str, str]:
    """
    Extract text from uploaded file based on its format.
    
    Args:
        file_content (bytes): Content of the uploaded file
        filename (str): Name of the uploaded file
        
    Returns:
        Tuple[str, str]: Extracted text content and file format
        
    Raises:
        UnsupportedFileFormatError: If file format is not supported
        FileSizeExceededError: If file size exceeds limit
        CorruptedFileError: If file is corrupted or unreadable
    """
    # Validate file format
    validate_file_format(filename)
    
    # Validate file size
    validate_file_size(len(file_content))
    
    # Determine file extension
    file_extension = '.' + filename.split('.')[-1].lower()
    
    try:
        if file_extension == '.txt':
            text_content = extract_text_from_txt(file_content)
        elif file_extension == '.docx':
            text_content = extract_text_from_docx(file_content)
        else:
            # This shouldn't happen due to validation, but just in case
            raise UnsupportedFileFormatError(f"Unsupported file format: {file_extension}")
        
        # Final validation - ensure we have meaningful content
        if not text_content or len(text_content.strip()) < 10:
            raise CorruptedFileError("Document contains insufficient readable content")
        
        logger.info(f"Successfully processed {filename} ({file_extension})")
        return text_content, file_extension
        
    except (UnsupportedFileFormatError, FileSizeExceededError, CorruptedFileError):
        # Re-raise our custom exceptions
        raise
    except Exception as e:
        logger.error(f"Unexpected error processing file {filename}: {str(e)}")
        raise CorruptedFileError(f"Unexpected error processing file: {str(e)}")